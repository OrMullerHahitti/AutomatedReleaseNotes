import asyncio
from typing import Any

from sqlalchemy.sql.schema import DefaultGenerator

from backend.app.models.base_service import BaseGenerator
from backend.app.models.models import LLMResponse
from backend.app.utils.prompts.prompts import default_prompt_format
from backend.app.utils.prompts.system_instructions import default_system_instructions
from docx import Document



class DefaultGenerator(BaseGenerator):
    ''' Default generator for release notes.
     This generator uses LangChain to generate release notes based on a list of work items.
     The release notes are generated using a structured JSON output that includes the title, doc name, and content.
    '''

    def __init__(self, llm: Any) -> None:
        super().__init__()
        self.llm = llm

    async def generate_doc(self, work_items: str,system_instructions: str=default_system_instructions,prompt_format:str=default_prompt_format) :
        """
           Creates a DOCX release note using LangChain and OpenAI.
           Automatically generates the title, doc name, and release note content from structured JSON output.

           :param system_instructions: Instructions or system prompt to guide the LLM.
           :param prompt_format: The format of the prompt to provide structured JSON output.
           :param work_items: A long string containing the list of work items.
           """

        # Initialize the LangChain LLM



        # Wrap the LLM with structured output
        llm_with_structured_output = self.llm.with_structured_output(LLMResponse)

        # Create a prompt template that specifies the output should be in JSON format
        prompt_template = PromptTemplate(
            input_variables=["system_instructions", "prompt_format", "work_items"],
            template=(
                "{system_instructions}\n\n"
                "{prompt_format}\n\n"
                "here is what you should use: List of work items:\n{work_items}\n"
            )
        )

        # Render the final prompt
        final_prompt = prompt_template.format(
            system_instructions=system_instructions,
            prompt_format=prompt_format,
            work_items=work_items
        )

        # Call the LLM to get the structured output
        response = llm_with_structured_output.invoke(final_prompt)

        # Create a new DOCX document
        doc = Document()

        # Add the title
        doc.add_heading(response.title, level=1)

        # Add the generated release note content
        doc.add_paragraph(response.content)

        # Save the document
        doc.save(response.doc_name if none)
        print(f"Release note saved to {response.doc_name}")


#TODO for next models, classification of work items and generation of release notes, more complex logic will be added



if __name__ == "__main__":
    from azure_authentication_client import authenticate_openai
    from langchain_openai import AzureChatOpenAI
    from langchain.prompts import PromptTemplate
    from backend.app.utils.useful_functions import format_work_items
    from backend.app.models.base_service import BaseGenerator
    from backend.app.services.azure_devops_services import AzureDevOpsService

    import os

    api_key = authenticate_openai().api_key
    if not os.environ.get("AZURE_OPENAI_API_KEY"):
        os.environ["AZURE_OPENAI_API_KEY"] = api_key

    # from backend.app.utils.config import summarize_prompt_text, format_prompt_text, release_notes_prompt_text

    api_key = authenticate_openai().api_key

    llm1 = AzureChatOpenAI(deployment_name="gpt-4o-deployment",
                          temperature=0.7,
                          azure_endpoint='https://function-app-open-ai-prod-apim.azure-api.net/proxy-api/')
    async def testing_one_two(language):
        test_service = AzureDevOpsService()
        work_items = await test_service.fetch_work_items("sprint 30")
        work_items= format_work_items(work_items)
        gen = DefaultGenerator(language)
        response = await gen.generate_doc(work_items)
    asyncio.run(testing_one_two(llm1))




