import asyncio
from abc import abstractmethod, ABC
from typing import List

from docx import Document

from langchain_core.output_parsers import PydanticOutputParser

from backend.app.models.models import WorkItem, TopicStructured, LLMResponse
from backend.app.services.llm_services.classifier import BaseClassifer
from backend.app.services.llm_services.llm_plugs.prompts import *
from backend.app.services.llm_services.llm_plugs import *
from backend.app.services.llm_services.llm_utils import generate_release_notes_paragraphs
from backend.app.utils.useful_functions import format_work_items, get_azure_llm
from backend.app.services.llm_services.llm_plugs.prompts import paragraph_examples as example_outputs
class BaseGenerator(ABC):
    """
    Abstract base class defining the 'generate_release_notes' interface.
    All concrete policies must implement this method.
    """
    def __init__(self,llm):
        self.llm=llm
        self.llm_response_output_parser = PydanticOutputParser(pydantic_object=LLMResponse)
        self.to_doc_template = PromptTemplate(
            template="""Given the following release note content:
            {content} format it given the parser given"""
        )
    async def generate(self, work_items: List[WorkItem]) -> LLMResponse:
        """
        Generate release notes from a list of WorkItems using the provided LLM.
        """
        work_items= format_work_items(work_items)
        docs= await self.generate_release_notes(llm=self.llm,work_items=work_items)
        release_notes = self.llm_response_output_parser.parse(docs['text'])
        doc = Document()

        # Add the title
        doc.add_heading(release_notes.title, level=1)

        # Add the generated release note content
        doc.add_paragraph(release_notes.content)

        # Save the document
        doc.save(f'{release_notes.doc_name}.docx' if release_notes.doc_name else "release_notes.docx")
        print(f"Release note saved to {release_notes.doc_name}")
        return release_notes

    @abstractmethod
    async def generate_release_notes(self,llm,  work_items: str) -> str:
        """
        Generate release notes from a list of WorkItems using the provided LLM.
        """
        pass

class BasicGenerator(BaseGenerator):
    ''' Default generator for release notes.
     This generator uses LangChain to generate release notes based on a list of work items.
     The release notes are generated using a structured JSON output that includes the title, doc name, and content.
    '''
    def __init__(self,llm):
        super().__init__(llm)


    async def generate_release_notes(self,llm, work_items: str) :
        """
           Creates a DOCX release note using LangChain and OpenAI.
           Automatically generates the title, doc name, and release note content from structured JSON output.

           :param system_instructions: Instructions or system prompt to guide the LLM.
           :param prompt_format: The format of the prompt to provide structured JSON output.
           :param work_items: A long string containing the list of work items.
           """
        classifier = BaseClassifer(llm, work_items)
        topics = await classifier.classify()
        rn_content = await generate_release_notes_paragraphs(llm,topic_data=topics,system_instructions=[system_insturctions.system_two],example_outputs=example_outputs)
        return rn_content


#TODO for next models, classification of work items and generation of release notes, more complex logic will be added



if __name__ == "__main__":
    from azure_authentication_client import authenticate_openai
    from langchain_openai import AzureChatOpenAI
    from langchain.prompts import PromptTemplate
    from backend.app.services.azure_devops_services import AzureDevOpsService

    from backend.app.utils.useful_functions import save_object_to_pickle,load_dict_from_pickle
    llm1 = get_azure_llm()
    async def testing_one_two(language):
        test_service = AzureDevOpsService()
        work_items = load_dict_from_pickle('3031.pickle')
        #work_items = await test_service.fetch_work_items_for_multiple_sprints(["Sprint 30", "Sprint 31"])
        gen = BasicGenerator(llm1)
        response = await gen.generate(work_items)
        print(response)
    asyncio.run(testing_one_two(llm1))
    print("hold")




