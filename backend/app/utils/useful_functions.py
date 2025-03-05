import os
from typing import List
import httpx
from azure_authentication_client import authenticate_openai
from langchain_openai import AzureChatOpenAI
from backend.app.models.models import WorkItem
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO


# <<<<<<< HEAD
# =======
from backend.app.utils.token_getter_llm import get_token_urlib
# >>>>>>> 20d549765888f3e22fe62d31f2ac15bcf9043a41


def format_work_items(work_items:List[WorkItem]):
    return "\n".join(f"- {item.title}: {item.description} (Type: {item.type}, State: {item.state} , {item.id})"
                     for item in work_items)


async def make_request(url: str, method: str = 'POST', headers: dict = None, data: dict = None):
    """
    Makes an asynchronous HTTP request using the specified method.

    Args:
        url (str): The URL to which the request is made.
        method (str): The HTTP method to use for the request ('GET' or 'POST'). Defaults to 'POST'.
        headers (dict, optional): A dictionary of HTTP headers to include in the request. Defaults to None.
        data (dict, optional): A dictionary of data to include in the request body for POST requests. Defaults to None.

    Returns:
        httpx.Response: The response object from the HTTP request.

    Raises:
        ValueError: If an unsupported HTTP method is specified.
        httpx.HTTPStatusError: If the HTTP request returns an unsuccessful status code.
    """
    async with httpx.AsyncClient() as client:
        if method == 'GET':
            response = await client.get(url, headers=headers, follow_redirects=False)
        elif method == 'POST':
            response = await client.post(url, headers=headers, json=data)
        else:
            raise ValueError(method)
        response.raise_for_status()
        return response


def get_azure_llm():
    api_key = get_token_urlib()
    if not os.environ.get("OPENAI_API_KEY"):
        os.environ["OPENAI_API_KEY"] = api_key
    return AzureChatOpenAI(
        deployment_name="gpt-4o-deployment",
        temperature=0.7,
        azure_endpoint="https://function-app-open-ai-prod-apim.azure-api.net/proxy-api/"
    )


def parse_html(html_string: str) -> str:
    soup = BeautifulSoup(html_string, 'html.parser')
    return soup.get_text()


def convert_text_to_docx(title: str , content: str) -> Document:

    """
        Converts the given text content into a Word document with a title and content.

        Args:
        - title (str): The title of the document.
        - content (str): The content to be included in the document.

        Returns:
        - Document: A python-docx Document object containing the title and content.
    """
    doc = Document()
    doc.add_heading(title, level=1) # Add the title
    doc.add_paragraph(content) # Add the generated release note content
    return doc
import pickle

def save_object_to_pickle(data, filename):
    """Saves a dictionary to a pickle file."""
    with open(filename, 'wb') as f:  # 'wb' for binary write
        pickle.dump(data, f)

def load_dict_from_pickle(filename):
    """Loads a dictionary from a pickle file."""
    with open(filename, 'rb') as f:  # 'rb' for binary read
        return pickle.load(f)

def convert_docx_to_bytes(doc: Document) -> bytes:
    """Converts a `Document` object to bytes."""
    byte_io = BytesIO()
    doc.save(byte_io)
    return byte_io.getvalue()
